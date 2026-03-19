"""文档解析器 - 读取目录下的所有文档并提取文本"""
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Dict, Any
import json

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    import markdown
except ImportError:
    markdown = None

@dataclass
class Page:
    """文档页面/章节"""
    page_num: int
    text: str
    tables: List[List[List[str]]] = field(default_factory=list)

@dataclass
class Document:
    """文档"""
    filename: str
    filepath: str
    pages: List[Page] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    def full_text(self) -> str:
        """获取完整文本"""
        return "\n\n".join(p.text for p in self.pages)

    def to_dict(self) -> dict:
        return {
            "filename": self.filename,
            "filepath": self.filepath,
            "metadata": self.metadata,
            "pages": [
                {"page_num": p.page_num, "text": p.text, "tables": p.tables}
                for p in self.pages
            ]
        }

class DocumentParser:
    """文档解析器 - 支持 PDF/DOCX/TXT/Markdown"""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".txt", ".md", ".markdown"}

    def __init__(self, doc_dir: str):
        self.doc_dir = Path(doc_dir)
        if not self.doc_dir.exists():
            raise FileNotFoundError(f"文档目录不存在: {doc_dir}")

    def parse_single(self, doc_path: str) -> Document:
        """解析单个文档"""
        path = Path(doc_path)
        if not path.exists():
            raise FileNotFoundError(f"文档文件不存在: {doc_path}")

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix == ".docx":
            return self._parse_docx(path)
        elif suffix == ".txt":
            return self._parse_txt(path)
        elif suffix in {".md", ".markdown"}:
            return self._parse_markdown(path)
        else:
            raise ValueError(f"不支持的格式: {suffix}")

    def _parse_pdf(self, path: Path) -> Document:
        """解析PDF"""
        if pdfplumber is None:
            raise ImportError("请安装 pdfplumber: pip install pdfplumber")

        doc = Document(filename=path.name, filepath=str(path))

        with pdfplumber.open(path) as pdf:
            doc.metadata = {
                "title": pdf.metadata.get("Title", ""),
                "author": pdf.metadata.get("Author", ""),
                "page_count": len(pdf.pages)
            }

            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                tables = page.extract_tables() or []
                doc.pages.append(Page(page_num=i, text=text, tables=tables))

        return doc

    def _parse_docx(self, path: Path) -> Document:
        """解析DOCX"""
        if docx is None:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(filename=path.name, filepath=str(path))

        doc_obj = docx.Document(path)
        doc.metadata = {
            "title": doc_obj.core_properties.title or "",
            "author": doc_obj.core_properties.author or "",
        }

        text_parts = []
        tables = []
        for para in doc_obj.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for i, table in enumerate(doc_obj.tables, 1):
            table_data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(table_data)
            text_parts.append(f"[表格 {i}]")

        doc.pages.append(Page(
            page_num=1,
            text="\n".join(text_parts),
            tables=tables
        ))

        return doc

    def _parse_txt(self, path: Path) -> Document:
        """解析TXT"""
        doc = Document(filename=path.name, filepath=str(path))
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()

        doc.pages.append(Page(page_num=1, text=text))
        return doc

    def _parse_markdown(self, path: Path) -> Document:
        """解析Markdown"""
        doc = Document(filename=path.name, filepath=str(path))

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        if markdown is not None:
            html = markdown.markdown(content)
            doc.metadata["html"] = html

        doc.pages.append(Page(page_num=1, text=content))
        return doc

    def parse_all(self, recursive: bool = False) -> List[Document]:
        """解析目录下所有支持的文档"""
        documents = []

        for suffix in self.SUPPORTED_FORMATS:
            pattern = f"**/*{suffix}" if recursive else f"*{suffix}"
            for doc_file in self.doc_dir.glob(pattern):
                try:
                    doc = self.parse_single(doc_file)
                    documents.append(doc)
                    print(f"✓ 解析完成: {doc.filename} ({len(doc.pages)} 页/节)")
                except Exception as e:
                    print(f"✗ 解析失败: {doc_file.name} - {e}")

        return documents

    def save_to_json(self, documents: List[Document], output_path: str):
        """保存解析结果到JSON"""
        data = [doc.to_dict() for doc in documents]
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"已保存到: {output_path}")
